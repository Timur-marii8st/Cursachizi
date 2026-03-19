"""Tests for ГОСТ-compliant document generation."""

import io

import pytest
from docx import Document

from backend.app.pipeline.formatter.docx_generator import DocxGenerator
from shared.schemas.pipeline import (
    BibliographyRegistry,
    Outline,
    OutlineChapter,
    SectionContent,
    Source,
)
from shared.schemas.template import GostTemplate, MarginConfig


@pytest.fixture
def generator() -> DocxGenerator:
    return DocxGenerator()


@pytest.fixture
def sample_outline() -> Outline:
    return Outline(
        title="Влияние цифровизации на управление персоналом",
        introduction_points=["Актуальность", "Цель исследования"],
        chapters=[
            OutlineChapter(
                number=1,
                title="Теоретические основы",
                subsections=["1.1 Понятия и определения", "1.2 Обзор литературы"],
                estimated_pages=10,
            ),
            OutlineChapter(
                number=2,
                title="Практический анализ",
                subsections=["2.1 Методология", "2.2 Результаты"],
                estimated_pages=12,
            ),
        ],
        conclusion_points=["Основные выводы"],
    )


@pytest.fixture
def sample_sections() -> list[SectionContent]:
    return [
        SectionContent(
            chapter_number=0,
            section_title="Введение",
            content="Данная курсовая работа посвящена исследованию влияния цифровизации.",
            word_count=10,
        ),
        SectionContent(
            chapter_number=1,
            section_title="1.1 Понятия и определения",
            content="Цифровизация — это процесс внедрения цифровых технологий [1].",
            word_count=10,
            citations=["1"],
        ),
        SectionContent(
            chapter_number=1,
            section_title="1.2 Обзор литературы",
            content="Многие исследователи изучали данную проблему [2].",
            word_count=8,
            citations=["2"],
        ),
        SectionContent(
            chapter_number=2,
            section_title="2.1 Методология",
            content="В исследовании применялись методы анализа и синтеза.",
            word_count=8,
        ),
        SectionContent(
            chapter_number=2,
            section_title="2.2 Результаты",
            content="Результаты исследования показали положительную динамику [3].",
            word_count=8,
            citations=["3"],
        ),
        SectionContent(
            chapter_number=99,
            section_title="Заключение",
            content="В ходе исследования были достигнуты поставленные цели.",
            word_count=8,
        ),
    ]


@pytest.fixture
def sample_bib_sources() -> list[Source]:
    return [
        Source(url="https://example.com/1", title="Иванов И.И. Цифровизация HR"),
        Source(url="https://example.com/2", title="Петров П.П. Управление персоналом"),
        Source(url="https://example.com/3", title="Сидоров С.С. Цифровая экономика"),
    ]


class TestDocxGenerator:
    def test_generates_valid_docx(
        self,
        generator: DocxGenerator,
        sample_outline: Outline,
        sample_sections: list[SectionContent],
        sample_bib_sources: list[Source],
    ) -> None:
        doc_bytes = generator.generate(
            outline=sample_outline,
            sections=sample_sections,
            sources=sample_bib_sources,
            university="Тестовый Университет",
            discipline="Менеджмент",
        )

        assert isinstance(doc_bytes, bytes)
        assert len(doc_bytes) > 0

        # Verify it's a valid docx by opening it
        doc = Document(io.BytesIO(doc_bytes))
        assert len(doc.paragraphs) > 0

    def test_contains_title_page(
        self,
        generator: DocxGenerator,
        sample_outline: Outline,
        sample_sections: list[SectionContent],
        sample_bib_sources: list[Source],
    ) -> None:
        doc_bytes = generator.generate(
            outline=sample_outline,
            sections=sample_sections,
            sources=sample_bib_sources,
            university="МГУ",
        )

        doc = Document(io.BytesIO(doc_bytes))
        text = "\n".join(p.text for p in doc.paragraphs)

        assert "МГУ" in text
        assert "КУРСОВАЯ РАБОТА" in text
        assert sample_outline.title in text

    def test_contains_all_sections(
        self,
        generator: DocxGenerator,
        sample_outline: Outline,
        sample_sections: list[SectionContent],
        sample_bib_sources: list[Source],
    ) -> None:
        doc_bytes = generator.generate(
            outline=sample_outline,
            sections=sample_sections,
            sources=sample_bib_sources,
        )

        doc = Document(io.BytesIO(doc_bytes))
        text = "\n".join(p.text for p in doc.paragraphs)

        assert "ВВЕДЕНИЕ" in text
        assert "ЗАКЛЮЧЕНИЕ" in text
        assert "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ" in text

    def test_bibliography_entries(
        self,
        generator: DocxGenerator,
        sample_outline: Outline,
        sample_sections: list[SectionContent],
        sample_bib_sources: list[Source],
    ) -> None:
        doc_bytes = generator.generate(
            outline=sample_outline,
            sections=sample_sections,
            sources=sample_bib_sources,
        )

        doc = Document(io.BytesIO(doc_bytes))
        text = "\n".join(p.text for p in doc.paragraphs)

        for source in sample_bib_sources:
            assert source.title in text

    def test_page_margins(
        self,
        sample_outline: Outline,
        sample_sections: list[SectionContent],
        sample_bib_sources: list[Source],
    ) -> None:
        """Verify ГОСТ margins: top=20, bottom=20, left=30, right=15."""
        generator = DocxGenerator()
        doc_bytes = generator.generate(
            outline=sample_outline,
            sections=sample_sections,
            sources=sample_bib_sources,
        )

        doc = Document(io.BytesIO(doc_bytes))
        section = doc.sections[0]

        # Margins in EMU (English Metric Units), 1mm = 36000 EMU
        mm_to_emu = 36000
        assert abs(section.top_margin - 20 * mm_to_emu) < mm_to_emu  # ±1mm tolerance
        assert abs(section.bottom_margin - 20 * mm_to_emu) < mm_to_emu
        assert abs(section.left_margin - 30 * mm_to_emu) < mm_to_emu
        assert abs(section.right_margin - 15 * mm_to_emu) < mm_to_emu

    def test_custom_template(
        self,
        sample_outline: Outline,
        sample_sections: list[SectionContent],
        sample_bib_sources: list[Source],
    ) -> None:
        """Test with custom margin configuration."""
        template = GostTemplate(
            margins=MarginConfig(top_mm=25, bottom_mm=25, left_mm=35, right_mm=10),
        )
        generator = DocxGenerator(template=template)

        doc_bytes = generator.generate(
            outline=sample_outline,
            sections=sample_sections,
            sources=sample_bib_sources,
        )

        doc = Document(io.BytesIO(doc_bytes))
        section = doc.sections[0]

        mm_to_emu = 36000
        assert abs(section.left_margin - 35 * mm_to_emu) < mm_to_emu

    def test_bibliography_from_registry(
        self,
        generator: DocxGenerator,
        sample_outline: Outline,
        sample_sections: list[SectionContent],
        sample_bib_sources: list[Source],
    ) -> None:
        """When BibliographyRegistry is provided, it should be used for bibliography."""
        registry = BibliographyRegistry.from_sources(sample_bib_sources)
        doc_bytes = generator.generate(
            outline=sample_outline,
            sections=sample_sections,
            sources=sample_bib_sources,
            bibliography=registry,
        )

        doc = Document(io.BytesIO(doc_bytes))
        text = "\n".join(p.text for p in doc.paragraphs)

        # All sources should appear in bibliography
        for source in sample_bib_sources:
            assert source.title in text
        # Should have [Электронный ресурс] for URL sources
        assert "[Электронный ресурс]" in text

    def test_bibliography_registry_preferred_over_extracted_refs(
        self,
        generator: DocxGenerator,
        sample_outline: Outline,
        sample_bib_sources: list[Source],
    ) -> None:
        """Registry sources should appear, not LLM-hallucinated ones."""
        # Section with a fake bibliography block that LLM might generate
        sections = [
            SectionContent(
                chapter_number=0,
                section_title="Введение",
                content="Текст введения [1].\n\n[1] Фейковый Автор. Фейковая Книга. — М., 2020.",
                word_count=10,
            ),
            SectionContent(
                chapter_number=99,
                section_title="Заключение",
                content="Текст заключения.",
                word_count=5,
            ),
        ]
        registry = BibliographyRegistry.from_sources(sample_bib_sources)
        doc_bytes = generator.generate(
            outline=Outline(title="Тест", chapters=[]),
            sections=sections,
            sources=sample_bib_sources,
            bibliography=registry,
        )

        doc = Document(io.BytesIO(doc_bytes))
        text = "\n".join(p.text for p in doc.paragraphs)

        # Real sources from registry should be in bibliography
        assert "Иванов И.И." in text
        # Fake source should NOT be in bibliography section
        # (it may still be in body text if not stripped, but bibliography is from registry)
        bib_start = text.find("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
        bib_text = text[bib_start:] if bib_start >= 0 else ""
        assert "Фейковый Автор" not in bib_text


class TestStripLeadingHeading:
    """Test _strip_leading_heading static method."""

    def test_strips_uppercase_heading(self):
        text = "ВВЕДЕНИЕ\nТекст введения начинается здесь."
        result = DocxGenerator._strip_leading_heading(text, "ВВЕДЕНИЕ")
        assert result == "Текст введения начинается здесь."

    def test_strips_heading_with_colon(self):
        text = "ВВЕДЕНИЕ: Текст введения."
        result = DocxGenerator._strip_leading_heading(text, "ВВЕДЕНИЕ")
        assert result == "Текст введения."

    def test_strips_heading_with_dash(self):
        text = "ЗАКЛЮЧЕНИЕ — В ходе исследования..."
        result = DocxGenerator._strip_leading_heading(text, "ЗАКЛЮЧЕНИЕ")
        assert result == "В ходе исследования..."

    def test_preserves_text_without_heading(self):
        text = "Текст без заголовка в начале."
        result = DocxGenerator._strip_leading_heading(text, "ВВЕДЕНИЕ")
        assert result == "Текст без заголовка в начале."

    def test_case_insensitive(self):
        text = "Введение\nТекст."
        result = DocxGenerator._strip_leading_heading(text, "ВВЕДЕНИЕ")
        assert result == "Текст."

    def test_strips_with_leading_whitespace(self):
        text = "  ВВЕДЕНИЕ\nТекст."
        result = DocxGenerator._strip_leading_heading(text, "ВВЕДЕНИЕ")
        assert result == "Текст."


class TestStripQuotes:
    """Test _strip_quotes static method."""

    def test_strips_guillemets(self):
        assert DocxGenerator._strip_quotes("«Менеджмент»") == "Менеджмент"

    def test_strips_double_quotes(self):
        assert DocxGenerator._strip_quotes('"Менеджмент"') == "Менеджмент"

    def test_preserves_unquoted(self):
        assert DocxGenerator._strip_quotes("Менеджмент") == "Менеджмент"

    def test_preserves_mismatched_quotes(self):
        assert DocxGenerator._strip_quotes("«Менеджмент") == "«Менеджмент"

    def test_strips_with_whitespace(self):
        assert DocxGenerator._strip_quotes("  «Менеджмент»  ") == "Менеджмент"


class TestCitationFixer:
    """Test citation_fixer integration with docx generation."""

    def test_strips_fake_bibliography_block(self):
        from backend.app.pipeline.writer.citation_fixer import fix_citations

        sources = [
            Source(url="https://real.com/1", title="Real Source One"),
            Source(url="https://real.com/2", title="Real Source Two"),
        ]
        registry = BibliographyRegistry.from_sources(sources)

        sections = [
            SectionContent(
                chapter_number=1,
                section_title="Test",
                content=(
                    "Some text [1] and [2] in body.\n\n"
                    "[1] Фейковый Автор. Фейковая Книга. — М., 2020.\n"
                    "[2] Другой Автор. Другая Книга. — СПб., 2021."
                ),
                word_count=10,
            ),
        ]

        result = fix_citations(sections, registry)
        assert len(result) == 1
        # Fake block should be stripped
        assert "Фейковый Автор" not in result[0].content
        # Inline citations should be present (possibly remapped)
        assert "[" in result[0].content

    def test_strips_bibliography_header(self):
        from backend.app.pipeline.writer.citation_fixer import fix_citations

        sources = [Source(url="https://example.com", title="Source")]
        registry = BibliographyRegistry.from_sources(sources)

        sections = [
            SectionContent(
                chapter_number=1,
                section_title="Test",
                content="Body text [1].\n\nБиблиографические ссылки:",
                word_count=5,
            ),
        ]

        result = fix_citations(sections, registry)
        assert "Библиографические ссылки" not in result[0].content

    def test_remaps_out_of_range_citations(self):
        from backend.app.pipeline.writer.citation_fixer import fix_citations

        sources = [
            Source(url="https://example.com/1", title="Source 1"),
            Source(url="https://example.com/2", title="Source 2"),
        ]
        registry = BibliographyRegistry.from_sources(sources)

        sections = [
            SectionContent(
                chapter_number=1,
                section_title="Test",
                content="According to [15] and [23], this is important.",
                word_count=8,
            ),
        ]

        result = fix_citations(sections, registry)
        # Citations should be remapped to valid range [1]-[2]
        import re
        cited = set(int(m) for m in re.findall(r"\[(\d+)\]", result[0].content))
        assert all(1 <= n <= 2 for n in cited)

    def test_no_block_no_crash(self):
        from backend.app.pipeline.writer.citation_fixer import fix_citations

        sources = [Source(url="https://example.com", title="Source")]
        registry = BibliographyRegistry.from_sources(sources)

        sections = [
            SectionContent(
                chapter_number=1,
                section_title="Test",
                content="Just some text [1] without bibliography block.",
                word_count=7,
            ),
        ]

        result = fix_citations(sections, registry)
        assert "[1]" in result[0].content


class TestFullPipelineFlow:
    """End-to-end tests simulating the orchestrator → fix_citations → docx_generator flow."""

    def test_full_flow_bibliography_appears_in_docx(self):
        """Simulate complete pipeline: build registry → fix citations → generate docx.

        This is the exact flow that the orchestrator follows:
        1. Build BibliographyRegistry from research sources
        2. Write sections (LLM output with fake bibliography blocks)
        3. fix_citations() strips fake blocks and remaps citations
        4. DocxGenerator.generate() receives fixed sections + registry
        5. Bibliography section in DOCX should contain REAL sources
        """
        from backend.app.pipeline.writer.citation_fixer import fix_citations

        # Real research sources
        sources = [
            Source(url="https://example.com/article1", title="Иванов И.И. Цифровая трансформация"),
            Source(url="https://example.com/article2", title="Петров П.П. Управление персоналом"),
            Source(url="https://example.com/article3", title="Smith J. Digital HR Management"),
        ]
        registry = BibliographyRegistry.from_sources(sources)

        # LLM output with fake bibliography blocks (simulating real LLM behavior)
        sections = [
            SectionContent(
                chapter_number=0,
                section_title="Введение",
                content="Актуальность темы подтверждается исследованиями [1] и [2].",
                word_count=8,
            ),
            SectionContent(
                chapter_number=1,
                section_title="1.1 Теоретические основы",
                content=(
                    "Согласно исследованиям [1], цифровизация меняет подходы к управлению [2]. "
                    "Многие авторы [3] отмечают значимость этих изменений.\n\n"
                    "Библиографические ссылки:\n"
                    "[1] Козлов А.А. Менеджмент в эпоху цифровизации. — М.: Наука, 2023.\n"
                    "[2] Brown T. HR Technology Trends. — NY: Wiley, 2022.\n"
                    "[3] Сидорова М.В. Кадровый потенциал. — СПб.: Питер, 2024."
                ),
                word_count=30,
            ),
            SectionContent(
                chapter_number=99,
                section_title="Заключение",
                content="В заключение отметим важность цифровизации.",
                word_count=7,
            ),
        ]

        outline = Outline(
            title="Цифровизация управления персоналом",
            chapters=[
                OutlineChapter(
                    number=1,
                    title="Теоретические основы",
                    subsections=["1.1 Теоретические основы"],
                ),
            ],
        )

        # Step 1: fix_citations (as orchestrator does in stage 4c)
        fixed_sections = fix_citations(sections, registry)

        # Verify fake block was stripped
        for s in fixed_sections:
            assert "Козлов А.А." not in s.content
            assert "Brown T." not in s.content
            assert "Библиографические ссылки" not in s.content

        # Step 2: generate docx (as orchestrator does in stage 5)
        generator = DocxGenerator()
        doc_bytes = generator.generate(
            outline=outline,
            sections=fixed_sections,
            sources=sources,
            bibliography=registry,
        )

        # Step 3: verify bibliography appears in output
        doc = Document(io.BytesIO(doc_bytes))
        full_text = "\n".join(p.text for p in doc.paragraphs)

        # Bibliography section must exist and contain REAL sources
        assert "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ" in full_text
        bib_start = full_text.find("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
        bib_text = full_text[bib_start:]

        # All 3 real sources must be in bibliography
        assert "Иванов И.И. Цифровая трансформация" in bib_text
        assert "Петров П.П. Управление персоналом" in bib_text
        assert "Smith J. Digital HR Management" in bib_text
        assert "[Электронный ресурс]" in bib_text

        # Fake sources must NOT be in bibliography
        assert "Козлов А.А." not in bib_text
        assert "Brown T. HR Technology" not in bib_text

    def test_full_flow_no_registry_fallback(self):
        """When no registry is provided, legacy extraction should work."""
        sources = [
            Source(url="https://example.com/1", title="Source One"),
        ]

        sections = [
            SectionContent(
                chapter_number=0,
                section_title="Введение",
                content="Text with [1] citation.\n\n[1] Some reference. — M., 2023.",
                word_count=8,
            ),
            SectionContent(
                chapter_number=99,
                section_title="Заключение",
                content="Conclusion text.",
                word_count=3,
            ),
        ]

        outline = Outline(
            title="Test",
            chapters=[],
        )

        # No registry — legacy path
        generator = DocxGenerator()
        doc_bytes = generator.generate(
            outline=outline,
            sections=sections,
            sources=sources,
            bibliography=None,
        )

        doc = Document(io.BytesIO(doc_bytes))
        full_text = "\n".join(p.text for p in doc.paragraphs)

        assert "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ" in full_text
        # When no registry is provided, references extracted from section text
        # take priority over raw sources
        assert "Some reference" in full_text


class TestRegistryPreservation:
    """Regression tests for bibliography registry preservation in DOCX output."""

    def test_registry_with_many_sources_survives_docx_generation(
        self,
        generator: DocxGenerator,
    ) -> None:
        sources = [
            Source(url=f"https://example.com/{i}", title=f"Source {i}")
            for i in range(1, 13)
        ]
        registry = BibliographyRegistry.from_sources(sources)

        outline = Outline(
            title="Registry preservation",
            chapters=[
                OutlineChapter(
                    number=1,
                    title="Chapter 1",
                    subsections=["1.1 Section"],
                    description="",
                ),
            ],
        )
        sections = [
            SectionContent(
                chapter_number=0,
                section_title="Введение",
                content="Intro text [1].\n\n[1] Fake bibliography entry.",
                word_count=4,
            ),
            SectionContent(
                chapter_number=1,
                section_title="1.1 Section",
                content="Body text [2].",
                word_count=3,
            ),
            SectionContent(
                chapter_number=99,
                section_title="Заключение",
                content="Conclusion text [3].",
                word_count=3,
            ),
        ]

        doc_bytes = generator.generate(
            outline=outline,
            sections=sections,
            sources=sources,
            bibliography=registry,
        )

        doc = Document(io.BytesIO(doc_bytes))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        bib_start = full_text.find("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
        assert bib_start >= 0
        bib_text = full_text[bib_start:]

        for source in sources:
            assert source.title in bib_text

        assert "Fake bibliography entry" not in bib_text


class TestEmptyBibliography:
    """Tests for empty bibliography edge case — should skip bibliography section."""

    def test_empty_bibliography_and_no_sources_skips_section(
        self,
        generator: DocxGenerator,
        sample_outline: Outline,
        sample_sections: list[SectionContent],
    ) -> None:
        """When bibliography is empty and sources list is empty,
        the bibliography section heading should NOT appear."""
        doc_bytes = generator.generate(
            outline=sample_outline,
            sections=sample_sections,
            sources=[],
            bibliography=BibliographyRegistry(entries=[]),
        )

        doc = Document(io.BytesIO(doc_bytes))
        text = "\n".join(p.text for p in doc.paragraphs)

        assert "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ" not in text

    def test_empty_bibliography_with_sources_uses_fallback(
        self,
        generator: DocxGenerator,
        sample_outline: Outline,
        sample_sections: list[SectionContent],
        sample_bib_sources: list[Source],
    ) -> None:
        """When bibliography has no entries but raw sources exist,
        fall back to formatting raw sources."""
        doc_bytes = generator.generate(
            outline=sample_outline,
            sections=sample_sections,
            sources=sample_bib_sources,
            bibliography=BibliographyRegistry(entries=[]),
        )

        doc = Document(io.BytesIO(doc_bytes))
        text = "\n".join(p.text for p in doc.paragraphs)

        assert "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ" in text
        assert "Иванов" in text

    def test_none_bibliography_and_no_sources_skips_section(
        self,
        generator: DocxGenerator,
        sample_outline: Outline,
        sample_sections: list[SectionContent],
    ) -> None:
        """When bibliography is None and sources list is empty,
        skip bibliography section."""
        doc_bytes = generator.generate(
            outline=sample_outline,
            sections=sample_sections,
            sources=[],
            bibliography=None,
        )

        doc = Document(io.BytesIO(doc_bytes))
        text = "\n".join(p.text for p in doc.paragraphs)

        assert "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ" not in text


class TestBibliographyWithScrapedAndUnscrapedSources:
    """Tests for the critical scenario where some sources failed to scrape
    but should still appear in the bibliography (title + URL are valid)."""

    def test_sources_without_fulltext_still_in_bibliography(self):
        """Sources that failed to scrape (no full_text) must still appear in bibliography."""
        from backend.app.pipeline.research.ranker import SourceRanker

        sources = [
            Source(
                url="https://scraped.com/article",
                title="Успешно скачанная статья",
                full_text="Длинный текст статьи " * 50,
                relevance_score=0.8,
            ),
            Source(
                url="https://blocked.com/page",
                title="Заблокированный сайт",
                full_text="",  # scraping failed
                relevance_score=0.5,
            ),
            Source(
                url="https://timeout.com/research",
                title="Таймаут при скачивании",
                full_text="Short",  # too short to be useful for writing
                relevance_score=0.6,
            ),
        ]

        # Ranker should preserve all 3 sources (they all have title + url)
        ranker = SourceRanker()
        ranked = ranker.rank_and_filter(sources)
        assert len(ranked) == 3

        # Build bibliography — all 3 sources must be present
        registry = BibliographyRegistry.from_sources(ranked)
        assert len(registry.entries) == 3

        # Generate docx with this registry
        generator = DocxGenerator()
        outline = Outline(
            title="Тестовая работа",
            chapters=[
                OutlineChapter(number=1, title="Глава 1", subsections=["1.1 Раздел"]),
            ],
        )
        sections = [
            SectionContent(
                chapter_number=0, section_title="Введение",
                content="Актуальность темы [1] и [2].", word_count=5,
            ),
            SectionContent(
                chapter_number=1, section_title="1.1 Раздел",
                content="Текст раздела [1] и [3].", word_count=5,
            ),
            SectionContent(
                chapter_number=99, section_title="Заключение",
                content="Выводы.", word_count=1,
            ),
        ]

        doc_bytes = generator.generate(
            outline=outline, sections=sections,
            sources=ranked, bibliography=registry,
        )

        doc = Document(io.BytesIO(doc_bytes))
        full_text = "\n".join(p.text for p in doc.paragraphs)

        # ALL three sources must appear in bibliography
        assert "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ" in full_text
        bib_start = full_text.find("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
        bib_text = full_text[bib_start:]

        assert "Успешно скачанная статья" in bib_text
        assert "Заблокированный сайт" in bib_text
        assert "Таймаут при скачивании" in bib_text
        assert "[Электронный ресурс]" in bib_text
        assert "дата обращения:" in bib_text

    def test_full_pipeline_sources_survive_all_stages(self):
        """Simulate the complete pipeline: research -> ranker -> registry -> citations -> docx.

        This is the most critical integration test: it verifies that sources
        collected during research make it all the way to the final document's
        bibliography section, even when some sources failed to scrape.
        """
        from backend.app.pipeline.research.ranker import SourceRanker
        from backend.app.pipeline.writer.citation_fixer import fix_citations

        # Step 1: Simulate research results (mix of scraped and unscraped)
        research_sources = [
            Source(
                url="https://cyberleninka.ru/article/1",
                title="Анализ современных подходов к менеджменту",
                full_text="Полный текст статьи с киберленинки " * 30,
                relevance_score=0.9,
                is_academic=True,
            ),
            Source(
                url="https://habr.com/article/2",
                title="Обзор технологий управления",
                full_text="Текст с хабра " * 20,
                relevance_score=0.7,
            ),
            Source(
                url="https://elibrary.ru/item/3",
                title="Кадровый потенциал организации",
                full_text="",  # blocked by paywall
                relevance_score=0.6,
                is_academic=True,
            ),
            Source(
                url="https://rbc.ru/article/4",
                title="Новости экономики и бизнеса",
                full_text="Кор",  # too short
                relevance_score=0.4,
            ),
        ]

        # Step 2: Ranker filters and ranks (should keep all 4)
        ranker = SourceRanker()
        ranked = ranker.rank_and_filter(research_sources)
        assert len(ranked) == 4, f"Expected 4 sources, got {len(ranked)}"

        # Step 3: Build bibliography registry
        registry = BibliographyRegistry.from_sources(ranked)
        assert len(registry.entries) == 4, f"Expected 4 entries, got {len(registry.entries)}"

        # Step 4: Simulate LLM writing with fake bibliography blocks
        sections = [
            SectionContent(
                chapter_number=0, section_title="Введение",
                content=(
                    "Актуальность темы управления подтверждается исследованиями [1] и [2]. "
                    "Кадровый потенциал рассмотрен в работе [3]."
                ),
                word_count=15,
            ),
            SectionContent(
                chapter_number=1, section_title="1.1 Теория менеджмента",
                content=(
                    "Современные подходы [1] включают цифровизацию [2]. "
                    "Экономические аспекты описаны в [4].\n\n"
                    "Список литературы:\n"
                    "[1] Выдуманный А.А. Книга. — М., 2023.\n"
                    "[2] Fake B.B. Another book. — NY, 2022."
                ),
                word_count=20,
            ),
            SectionContent(
                chapter_number=99, section_title="Заключение",
                content="Работа показала важность управления [1].",
                word_count=5,
            ),
        ]

        # Step 5: Fix citations (strips fake blocks, remaps numbers)
        fixed = fix_citations(sections, registry)

        # Verify fake blocks stripped
        for s in fixed:
            assert "Выдуманный" not in s.content
            assert "Fake B.B." not in s.content
            assert "Список литературы" not in s.content

        # Step 6: Generate DOCX
        generator = DocxGenerator()
        outline = Outline(
            title="Управление персоналом в цифровую эпоху",
            chapters=[
                OutlineChapter(
                    number=1, title="Теория менеджмента",
                    subsections=["1.1 Теория менеджмента"],
                ),
            ],
        )

        doc_bytes = generator.generate(
            outline=outline, sections=fixed,
            sources=ranked, bibliography=registry,
        )

        # Step 7: Parse and verify final document
        doc = Document(io.BytesIO(doc_bytes))
        full_text = "\n".join(p.text for p in doc.paragraphs)

        # Bibliography section must exist
        assert "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ" in full_text, \
            "Bibliography section missing from final document!"

        bib_start = full_text.find("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
        bib_text = full_text[bib_start:]

        # ALL 4 real research sources must be in bibliography
        assert "Анализ современных подходов к менеджменту" in bib_text
        assert "Обзор технологий управления" in bib_text
        assert "Кадровый потенциал организации" in bib_text
        assert "Новости экономики и бизнеса" in bib_text

        # GOST formatting checks
        assert "[Электронный ресурс]" in bib_text
        assert "дата обращения:" in bib_text

        # Fake sources must NOT be present
        assert "Выдуманный" not in bib_text
        assert "Fake B.B." not in bib_text
