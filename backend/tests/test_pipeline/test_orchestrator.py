"""Tests for the pipeline orchestrator."""

import io
import json

import pytest
from docx import Document

from backend.app.pipeline.orchestrator import PipelineOrchestrator, StageCallback
from backend.app.testing import MockLLMProvider, MockSearchProvider
from shared.schemas.job import WorkType
from shared.schemas.pipeline import PipelineConfig, Source


@pytest.fixture
def search_with_results(mock_search: MockSearchProvider) -> MockSearchProvider:
    """Search provider pre-loaded with sample results."""
    mock_search.set_results([
        Source(
            url="https://example.com/1",
            title="Source 1",
            snippet="Snippet 1",
            full_text="Full text about the topic with enough content " * 20,
            relevance_score=0.8,
        ),
        Source(
            url="https://example.com/2",
            title="Source 2",
            snippet="Snippet 2",
            full_text="Another full text source with relevant information " * 15,
            relevance_score=0.7,
        ),
    ])
    return mock_search


class TrackingCallback(StageCallback):
    """Tracks stage transitions for testing."""

    def __init__(self) -> None:
        self.events: list[tuple[str, str, str]] = []

    async def on_stage_start(self, stage: str, message: str = "") -> None:
        self.events.append(("start", stage, message))

    async def on_stage_progress(
        self, stage: str, progress_pct: int, message: str = ""
    ) -> None:
        self.events.append(("progress", stage, message))

    async def on_stage_complete(self, stage: str, message: str = "") -> None:
        self.events.append(("complete", stage, message))


class TestPipelineOrchestrator:
    async def test_large_coursework_requires_minimum_source_pool(
        self,
        mock_llm: MockLLMProvider,
        search_with_results: MockSearchProvider,
    ) -> None:
        """30-page coursework should fail fast when research yields too few sources."""
        mock_llm.set_response(json.dumps({"queries": ["query1"]}))

        orchestrator = PipelineOrchestrator(
            llm=mock_llm,
            search=search_with_results,
        )

        with pytest.raises(RuntimeError, match="Insufficient research sources"):
            await orchestrator.run(
                topic="Large coursework topic",
                discipline="Test",
                page_count=30,
                config=PipelineConfig(enable_fact_check=False),
            )

    async def test_full_pipeline_execution(
        self,
        mock_llm: MockLLMProvider,
        search_with_results: MockSearchProvider,
    ) -> None:
        """Smoke test: full pipeline runs without error with mocks."""
        # Set up LLM responses for each stage:
        # 1. Query expansion
        # 2. Outline generation
        # 3+ Section writing (intro, sections, conclusion)
        # N+ Claim extraction per chapter
        # N+ Fact check verdicts

        query_response = json.dumps({"queries": ["query1", "query2"]})
        outline_response = json.dumps({
            "title": "Test Paper Title",
            "introduction_points": ["Point 1"],
            "chapters": [
                {
                    "number": 1,
                    "title": "Chapter 1",
                    "subsections": ["1.1 Section"],
                    "description": "Description",
                    "estimated_pages": 10,
                }
            ],
            "conclusion_points": ["Conclusion"],
        })
        # Include citation markers so body sections pass quality evaluation
        # and do not trigger LLM rewrites.
        section_text = "This is generated section content [1] and more [2]. " * 40
        claims_response = json.dumps({
            "claims": [
                {"claim_text": "Test claim", "source_section": "1.1"}
            ]
        })
        verdict_response = (
            "VERDICT: supported\n"
            "CONFIDENCE: 0.8\n"
            "EVIDENCE: Confirmed by sources.\n"
            "CORRECTION: нет"
        )

        # The pipeline will call LLM multiple times:
        # query_expand, outline, intro, section, conclusion, claim_extract, verdict,
        # plus padding for optional section rewrites and intro/conclusion validation.
        mock_llm.set_responses([
            query_response,    # query expansion
            outline_response,  # outline
            section_text,      # introduction
            section_text,      # section 1.1
            section_text,      # conclusion
            claims_response,   # claim extraction for section 1.1
            verdict_response,  # fact check verdict
            section_text,      # section rewrite padding (if triggered)
            section_text,      # intro_conclusion_validator fix_introduction (if triggered)
            section_text,      # intro_conclusion_validator fix_conclusion (if triggered)
        ])

        config = PipelineConfig(
            max_search_results=5,
            max_sources=5,
            max_claims_per_chapter=2,
            enable_fact_check=True,
        )

        orchestrator = PipelineOrchestrator(
            llm=mock_llm,
            search=search_with_results,
        )

        callback = TrackingCallback()
        result = await orchestrator.run(
            topic="Test topic",
            discipline="Test",
            page_count=20,
            config=config,
            callback=callback,
        )

        assert result.error is None
        assert result.outline is not None
        assert result.research is not None
        assert len(result.sections) > 0
        assert result.document_bytes is not None
        assert len(result.document_bytes) > 0
        assert result.completed_at is not None

        # Verify callback received stage events
        stage_names = [e[1] for e in callback.events if e[0] == "start"]
        assert "researching" in stage_names
        assert "outlining" in stage_names
        assert "writing" in stage_names
        assert "formatting" in stage_names

    async def test_pipeline_without_fact_check(
        self,
        mock_llm: MockLLMProvider,
        search_with_results: MockSearchProvider,
    ) -> None:
        query_response = json.dumps({"queries": ["q1"]})
        outline_response = json.dumps({
            "title": "Test",
            "introduction_points": [],
            "chapters": [
                {
                    "number": 1,
                    "title": "Ch1",
                    "subsections": ["1.1 S"],
                    "description": "D",
                    "estimated_pages": 5,
                }
            ],
            "conclusion_points": [],
        })
        # Include citation markers so the body section passes quality evaluation
        # and no LLM rewrite is triggered.
        section_text = "Content [1] content [2] content. " * 30

        mock_llm.set_responses([
            query_response,
            outline_response,
            section_text,  # intro
            section_text,  # 1.1
            section_text,  # conclusion
            section_text,  # intro_conclusion_validator fix_introduction (if triggered)
            section_text,  # intro_conclusion_validator fix_conclusion (if triggered)
        ])

        config = PipelineConfig(
            enable_fact_check=False,
            enable_section_rewrite=False,
            enable_coherence_check=False,
            max_search_results=5,
        )

        orchestrator = PipelineOrchestrator(
            llm=mock_llm, search=search_with_results
        )

        result = await orchestrator.run(
            topic="Test",
            page_count=20,
            config=config,
        )

        assert result.error is None
        assert result.fact_check is None
        assert result.document_bytes is not None

    async def test_pipeline_article_mode(
        self,
        mock_llm: MockLLMProvider,
        search_with_results: MockSearchProvider,
    ) -> None:
        """Test pipeline execution in article mode."""
        query_response = json.dumps({"queries": ["q1"]})
        outline_response = json.dumps({
            "title": "Анализ цифровой трансформации",
            "abstract_points": ["Цель исследования", "Основные результаты"],
            "keywords": ["цифровизация", "трансформация"],
            "introduction_points": ["Актуальность"],
            "sections": [
                {
                    "number": 1,
                    "title": "Теоретические основы",
                    "description": "Обзор литературы",
                    "estimated_pages": 3,
                },
                {
                    "number": 2,
                    "title": "Результаты и обсуждение",
                    "description": "Анализ данных",
                    "estimated_pages": 3,
                },
            ],
            "conclusion_points": ["Выводы"],
        })
        section_text = "Текст раздела научной статьи. " * 30

        mock_llm.set_responses([
            query_response,    # query expansion
            outline_response,  # outline
            section_text,      # abstract
            section_text,      # introduction
            section_text,      # section 1
            section_text,      # section 2
            section_text,      # conclusion
        ])

        config = PipelineConfig(
            enable_fact_check=False,
            max_search_results=5,
            enable_section_rewrite=False,
            enable_coherence_check=False,
        )

        orchestrator = PipelineOrchestrator(
            llm=mock_llm, search=search_with_results
        )

        callback = TrackingCallback()
        result = await orchestrator.run(
            topic="Анализ цифровой трансформации",
            discipline="Менеджмент",
            page_count=10,
            work_type=WorkType.ARTICLE,
            config=config,
            callback=callback,
        )

        assert result.error is None
        assert result.outline is not None
        assert result.outline.keywords == ["цифровизация", "трансформация"]
        assert len(result.sections) > 0
        assert result.document_bytes is not None
        assert len(result.document_bytes) > 0

        # Verify article-specific sections exist
        section_titles = [s.section_title for s in result.sections]
        assert "Аннотация" in section_titles
        assert "Введение" in section_titles
        assert "Заключение" in section_titles

        # Verify the generated docx does NOT contain "КУРСОВАЯ РАБОТА"
        doc = Document(io.BytesIO(result.document_bytes))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "КУРСОВАЯ РАБОТА" not in full_text, (
            "Article document must NOT contain 'КУРСОВАЯ РАБОТА'"
        )

    async def test_pipeline_coursework_has_correct_label(
        self,
        mock_llm: MockLLMProvider,
        search_with_results: MockSearchProvider,
    ) -> None:
        """Test coursework pipeline produces 'КУРСОВАЯ РАБОТА' label."""
        query_response = json.dumps({"queries": ["q1"]})
        outline_response = json.dumps({
            "title": "Тестовая курсовая",
            "introduction_points": [],
            "chapters": [
                {
                    "number": 1,
                    "title": "Глава 1",
                    "subsections": ["1.1 Раздел"],
                    "description": "Описание",
                    "estimated_pages": 5,
                }
            ],
            "conclusion_points": [],
        })
        section_text = "Текст раздела курсовой работы [1]. " * 30

        mock_llm.set_responses([
            query_response,
            outline_response,
            section_text,  # intro
            section_text,  # 1.1
            section_text,  # conclusion
            section_text,  # padding
            section_text,  # padding
        ])

        config = PipelineConfig(
            enable_fact_check=False,
            max_search_results=5,
            enable_section_rewrite=False,
            enable_coherence_check=False,
        )

        orchestrator = PipelineOrchestrator(
            llm=mock_llm, search=search_with_results
        )

        result = await orchestrator.run(
            topic="Тестовая курсовая",
            work_type=WorkType.COURSEWORK,
            page_count=20,
            config=config,
        )

        assert result.error is None
        assert result.document_bytes is not None

        doc = Document(io.BytesIO(result.document_bytes))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "КУРСОВАЯ РАБОТА" in full_text, (
            "Coursework document must contain 'КУРСОВАЯ РАБОТА'"
        )

    async def test_pipeline_article_has_bibliography(
        self,
        mock_llm: MockLLMProvider,
        search_with_results: MockSearchProvider,
    ) -> None:
        """Test that article pipeline produces bibliography when sources exist."""
        query_response = json.dumps({"queries": ["q1"]})
        outline_response = json.dumps({
            "title": "Тестовая статья",
            "abstract_points": ["Точка 1"],
            "keywords": ["тест"],
            "introduction_points": ["Актуальность"],
            "sections": [
                {
                    "number": 1,
                    "title": "Раздел 1",
                    "description": "Описание",
                    "estimated_pages": 3,
                },
            ],
            "conclusion_points": ["Вывод"],
        })
        section_text = "Текст с ссылками [1] и [2]. " * 30

        mock_llm.set_responses([
            query_response,
            outline_response,
            section_text,  # abstract
            section_text,  # intro
            section_text,  # section 1
            section_text,  # conclusion
        ])

        config = PipelineConfig(
            enable_fact_check=False,
            max_search_results=5,
            enable_section_rewrite=False,
            enable_coherence_check=False,
        )

        orchestrator = PipelineOrchestrator(
            llm=mock_llm, search=search_with_results
        )

        result = await orchestrator.run(
            topic="Тестовая статья",
            work_type=WorkType.ARTICLE,
            config=config,
        )

        assert result.error is None
        assert result.document_bytes is not None

        # Check that research found sources
        assert len(result.research.sources) > 0, "Research should find sources"
        assert result.bibliography is not None
        assert len(result.bibliography.entries) > 0, "Bibliography should have entries"

        doc = Document(io.BytesIO(result.document_bytes))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "СПИСОК ЛИТЕРАТУРЫ" in full_text, (
            "Article should contain bibliography section"
        )
        # Check at least one source appears
        assert "Source 1" in full_text or "Source 2" in full_text or "example.com" in full_text, (
            "Bibliography should contain actual sources"
        )
