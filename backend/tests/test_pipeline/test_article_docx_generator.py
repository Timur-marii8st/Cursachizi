"""Tests for ГОСТ-compliant article document generation."""

import io

import pytest
from docx import Document

from backend.app.pipeline.formatter.article_docx_generator import ArticleDocxGenerator
from shared.schemas.pipeline import (
    BibliographyRegistry,
    Outline,
    OutlineChapter,
    SectionContent,
    Source,
)
from shared.schemas.template import GostTemplate


@pytest.fixture
def generator() -> ArticleDocxGenerator:
    return ArticleDocxGenerator()


@pytest.fixture
def sample_outline() -> Outline:
    return Outline(
        title="Методы машинного обучения в задачах классификации текстов",
        introduction_points=["Актуальность темы", "Цель и задачи исследования"],
        chapters=[
            OutlineChapter(
                number=1,
                title="Теоретические основы классификации текстов",
                subsections=[],
                estimated_pages=5,
            ),
            OutlineChapter(
                number=2,
                title="Экспериментальное сравнение моделей",
                subsections=[],
                estimated_pages=7,
            ),
        ],
        conclusion_points=["Основные результаты", "Направления дальнейших исследований"],
        keywords=["машинное обучение", "классификация текстов", "NLP", "нейронные сети"],
        abstract_points=["Обзор методов", "Результаты экспериментов"],
    )


@pytest.fixture
def sample_sections() -> list[SectionContent]:
    return [
        SectionContent(
            chapter_number=0,
            section_title="Аннотация",
            content=(
                "В статье рассматриваются современные методы машинного обучения, "
                "применяемые для автоматической классификации текстовых документов."
            ),
            word_count=12,
        ),
        SectionContent(
            chapter_number=0,
            section_title="Введение",
            content=(
                "Задача автоматической классификации текстов является одной из "
                "ключевых задач обработки естественного языка [1]."
            ),
            word_count=14,
            citations=["1"],
        ),
        SectionContent(
            chapter_number=1,
            section_title="Теоретические основы классификации текстов",
            content=(
                "Классификация текстов представляет собой процесс отнесения "
                "текстового документа к одной или нескольким предопределённым "
                "категориям [2]."
            ),
            word_count=15,
            citations=["2"],
        ),
        SectionContent(
            chapter_number=2,
            section_title="Экспериментальное сравнение моделей",
            content=(
                "Для проведения эксперимента были выбраны три модели: "
                "логистическая регрессия, случайный лес и BERT [3]."
            ),
            word_count=13,
            citations=["3"],
        ),
        SectionContent(
            chapter_number=99,
            section_title="Заключение",
            content=(
                "В результате проведённого исследования установлено, что модели "
                "на основе трансформеров демонстрируют наилучшее качество классификации."
            ),
            word_count=14,
        ),
    ]


@pytest.fixture
def sample_sources() -> list[Source]:
    return [
        Source(url="https://example.com/nlp", title="Иванов А.А. Обработка естественного языка"),
        Source(
            url="https://example.com/classification",
            title="Петров Б.Б. Классификация текстов",
        ),
        Source(url="https://example.com/bert", title="Сидоров В.В. Модели трансформеров"),
    ]


def _generate_and_parse(
    generator: ArticleDocxGenerator,
    outline: Outline,
    sections: list[SectionContent],
    sources: list[Source],
    **kwargs: str,
) -> tuple[bytes, list[str]]:
    """Helper: generate docx bytes and extract paragraph texts."""
    result = generator.generate(
        outline=outline,
        sections=sections,
        sources=sources,
        **kwargs,
    )
    doc = Document(io.BytesIO(result))
    texts = [p.text for p in doc.paragraphs]
    return result, texts


class TestArticleDocxGenerator:
    def test_basic_generation_produces_valid_bytes(
        self,
        generator: ArticleDocxGenerator,
        sample_outline: Outline,
        sample_sections: list[SectionContent],
        sample_sources: list[Source],
    ) -> None:
        result = generator.generate(
            outline=sample_outline,
            sections=sample_sections,
            sources=sample_sources,
        )

        assert isinstance(result, bytes)
        assert len(result) > 0

        # Verify the bytes form a valid docx
        doc = Document(io.BytesIO(result))
        assert len(doc.paragraphs) > 0

    def test_article_structure_contains_required_elements(
        self,
        generator: ArticleDocxGenerator,
        sample_outline: Outline,
        sample_sections: list[SectionContent],
        sample_sources: list[Source],
    ) -> None:
        """Article must have: title, abstract, keywords, sections, conclusion, bibliography."""
        _, texts = _generate_and_parse(generator, sample_outline, sample_sections, sample_sources)
        full_text = "\n".join(texts)

        # Title (uppercased)
        assert sample_outline.title.upper() in full_text

        # Abstract label
        assert "Аннотация." in full_text

        # Keywords label
        assert "Ключевые слова:" in full_text

        # Introduction
        assert "Введение" in full_text

        # Chapter headings (flat, no sub-chapters)
        assert "1. Теоретические основы классификации текстов" in full_text
        assert "2. Экспериментальное сравнение моделей" in full_text

        # Conclusion
        assert "Заключение" in full_text

        # Bibliography heading
        assert "СПИСОК ЛИТЕРАТУРЫ" in full_text

    def test_article_has_no_coursework_markers(
        self,
        generator: ArticleDocxGenerator,
        sample_outline: Outline,
        sample_sections: list[SectionContent],
        sample_sources: list[Source],
    ) -> None:
        """Article must NOT contain coursework-specific labels."""
        _, texts = _generate_and_parse(generator, sample_outline, sample_sections, sample_sources)
        full_text = "\n".join(texts)

        assert "КУРСОВАЯ РАБОТА" not in full_text
        assert "СОДЕРЖАНИЕ" not in full_text

    def test_author_info_included_when_provided(
        self,
        generator: ArticleDocxGenerator,
        sample_outline: Outline,
        sample_sections: list[SectionContent],
        sample_sources: list[Source],
    ) -> None:
        _, texts = _generate_and_parse(
            generator,
            sample_outline,
            sample_sections,
            sample_sources,
            author="Козлов Д.Д.",
            university="Московский государственный университет",
        )
        full_text = "\n".join(texts)

        assert "Козлов Д.Д." in full_text
        assert "Московский государственный университет" in full_text

    def test_author_info_absent_when_not_provided(
        self,
        generator: ArticleDocxGenerator,
        sample_outline: Outline,
        sample_sections: list[SectionContent],
        sample_sources: list[Source],
    ) -> None:
        _, texts = _generate_and_parse(
            generator, sample_outline, sample_sections, sample_sources
        )
        full_text = "\n".join(texts)

        # No author-specific paragraph should appear; only content-related text
        # The title, abstract, keywords, sections, and bibliography are expected
        assert "Козлов" not in full_text

    def test_keywords_rendered_from_outline(
        self,
        generator: ArticleDocxGenerator,
        sample_outline: Outline,
        sample_sections: list[SectionContent],
        sample_sources: list[Source],
    ) -> None:
        _, texts = _generate_and_parse(generator, sample_outline, sample_sections, sample_sources)
        full_text = "\n".join(texts)

        for keyword in sample_outline.keywords:
            assert keyword in full_text, f"Keyword '{keyword}' not found in document"

        # Keywords should be comma-separated in a single paragraph
        keywords_para = [t for t in texts if "Ключевые слова:" in t]
        assert len(keywords_para) == 1
        kw_text = keywords_para[0]
        assert "машинное обучение" in kw_text
        assert "NLP" in kw_text

    def test_keywords_absent_when_empty(
        self,
        generator: ArticleDocxGenerator,
        sample_sections: list[SectionContent],
        sample_sources: list[Source],
    ) -> None:
        outline_no_kw = Outline(
            title="Статья без ключевых слов",
            chapters=[],
            keywords=[],
        )
        _, texts = _generate_and_parse(
            generator, outline_no_kw, sample_sections, sample_sources
        )
        full_text = "\n".join(texts)

        assert "Ключевые слова:" not in full_text

    def test_clean_text_removes_markdown_bold(self) -> None:
        cleaned = ArticleDocxGenerator._clean_text("Это **жирный** текст")
        assert "**" not in cleaned
        assert "жирный" in cleaned

    def test_clean_text_removes_markdown_italic(self) -> None:
        cleaned = ArticleDocxGenerator._clean_text("Это *курсивный* текст")
        assert cleaned == "Это курсивный текст"

    def test_clean_text_removes_inline_code(self) -> None:
        cleaned = ArticleDocxGenerator._clean_text("Используется `pandas` для анализа")
        assert "`" not in cleaned
        assert "pandas" in cleaned

    def test_clean_text_removes_markdown_links(self) -> None:
        cleaned = ArticleDocxGenerator._clean_text("Смотри [документацию](https://example.com)")
        assert "[" not in cleaned
        assert "https://example.com" not in cleaned
        assert "документацию" in cleaned

    def test_clean_text_removes_heading_markers(self) -> None:
        cleaned = ArticleDocxGenerator._clean_text("## Заголовок второго уровня")
        assert cleaned.strip() == "Заголовок второго уровня"

    def test_clean_text_removes_html_entities(self) -> None:
        cleaned = ArticleDocxGenerator._clean_text("Текст&nbsp;с&amp;сущностями")
        assert "&nbsp;" not in cleaned
        assert "&amp;" not in cleaned

    def test_clean_text_removes_list_markers(self) -> None:
        cleaned = ArticleDocxGenerator._clean_text("- Элемент списка\n* Другой элемент")
        assert "- " not in cleaned
        assert "* " not in cleaned
        assert "Элемент списка" in cleaned

    def test_bibliography_entries_present(
        self,
        generator: ArticleDocxGenerator,
        sample_outline: Outline,
        sample_sections: list[SectionContent],
        sample_sources: list[Source],
    ) -> None:
        _, texts = _generate_and_parse(generator, sample_outline, sample_sections, sample_sources)
        full_text = "\n".join(texts)

        for source in sample_sources:
            assert source.title in full_text, f"Source '{source.title}' not in bibliography"
            assert source.url in full_text, f"URL '{source.url}' not in bibliography"

    def test_bibliography_entries_numbered(
        self,
        generator: ArticleDocxGenerator,
        sample_outline: Outline,
        sample_sections: list[SectionContent],
        sample_sources: list[Source],
    ) -> None:
        _, texts = _generate_and_parse(generator, sample_outline, sample_sections, sample_sources)
        bib_texts = [t for t in texts if "Иванов" in t or "Петров" in t or "Сидоров" in t]

        assert any(t.startswith("1.") for t in bib_texts)
        assert any(t.startswith("2.") for t in bib_texts)
        assert any(t.startswith("3.") for t in bib_texts)

    def test_custom_template_margins(
        self,
        sample_outline: Outline,
        sample_sections: list[SectionContent],
        sample_sources: list[Source],
    ) -> None:
        from shared.schemas.template import MarginConfig

        template = GostTemplate(
            margins=MarginConfig(top_mm=25, bottom_mm=25, left_mm=35, right_mm=10),
        )
        gen = ArticleDocxGenerator(template=template)

        result = gen.generate(
            outline=sample_outline,
            sections=sample_sections,
            sources=sample_sources,
        )

        doc = Document(io.BytesIO(result))
        section = doc.sections[0]

        mm_to_emu = 36000
        assert abs(section.left_margin - 35 * mm_to_emu) < mm_to_emu
        assert abs(section.right_margin - 10 * mm_to_emu) < mm_to_emu
        assert abs(section.top_margin - 25 * mm_to_emu) < mm_to_emu
        assert abs(section.bottom_margin - 25 * mm_to_emu) < mm_to_emu


class TestArticleEmptyBibliography:
    """Tests for empty bibliography edge case in article generation."""

    def test_empty_bibliography_and_no_sources_skips_section(
        self,
        generator: ArticleDocxGenerator,
        sample_outline: Outline,
        sample_sections: list[SectionContent],
    ) -> None:
        """Bibliography section should be skipped when nothing to render."""
        doc_bytes = generator.generate(
            outline=sample_outline,
            sections=sample_sections,
            sources=[],
            bibliography=BibliographyRegistry(entries=[]),
        )

        doc = Document(io.BytesIO(doc_bytes))
        text = "\n".join(p.text for p in doc.paragraphs)

        assert "СПИСОК ЛИТЕРАТУРЫ" not in text

    def test_empty_bibliography_with_sources_uses_fallback(
        self,
        generator: ArticleDocxGenerator,
        sample_outline: Outline,
        sample_sections: list[SectionContent],
        sample_sources: list[Source],
    ) -> None:
        """When bibliography empty but raw sources exist, use fallback."""
        doc_bytes = generator.generate(
            outline=sample_outline,
            sections=sample_sections,
            sources=sample_sources,
            bibliography=BibliographyRegistry(entries=[]),
        )

        doc = Document(io.BytesIO(doc_bytes))
        text = "\n".join(p.text for p in doc.paragraphs)

        assert "СПИСОК ЛИТЕРАТУРЫ" in text

    def test_none_bibliography_and_no_sources_skips_section(
        self,
        generator: ArticleDocxGenerator,
        sample_outline: Outline,
        sample_sections: list[SectionContent],
    ) -> None:
        """When bibliography is None and no sources, skip section."""
        doc_bytes = generator.generate(
            outline=sample_outline,
            sections=sample_sections,
            sources=[],
            bibliography=None,
        )

        doc = Document(io.BytesIO(doc_bytes))
        text = "\n".join(p.text for p in doc.paragraphs)

        assert "СПИСОК ЛИТЕРАТУРЫ" not in text
