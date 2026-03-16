"""ГОСТ-compliant .docx document generator using python-docx."""

import io
import re

import structlog
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm, Pt, RGBColor

from backend.app.pipeline.formatter.reference_extractor import (
    extract_and_renumber_references,
)
from shared.schemas.pipeline import (
    BibliographyRegistry,
    FactCheckResult,
    Outline,
    SectionContent,
    Source,
)
from shared.schemas.template import GostTemplate

logger = structlog.get_logger()

# Alignment mapping
ALIGNMENT_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


class DocxGenerator:
    """Generates ГОСТ-compliant .docx files from coursework content."""

    def __init__(self, template: GostTemplate | None = None) -> None:
        self._template = template or GostTemplate()

    def generate(
        self,
        outline: Outline,
        sections: list[SectionContent],
        sources: list[Source],
        fact_check: FactCheckResult | None = None,
        university: str = "",
        discipline: str = "",
        author: str = "",
        bibliography: BibliographyRegistry | None = None,
    ) -> bytes:
        """Generate a complete .docx document.

        Args:
            outline: Coursework outline.
            sections: All written sections in order.
            sources: Research sources for bibliography.
            fact_check: Optional fact-check results (for annotation).
            university: University name for title page.
            discipline: Discipline for title page.
            author: Author name for title page.
            bibliography: Unified bibliography registry from real sources.

        Returns:
            Bytes of the generated .docx file.
        """
        # Citation fixing is done by the orchestrator (fix_citations) before
        # reaching the generator. Here we only handle the legacy path when no
        # bibliography registry is provided (e.g. direct generator usage).
        logger.info(
            "docx_generate_start",
            has_bibliography=bibliography is not None,
            bibliography_entries=len(bibliography.entries) if bibliography else 0,
            sections_count=len(sections),
            sources_count=len(sources),
        )
        if not (bibliography and bibliography.entries):
            ref_result = extract_and_renumber_references(sections)
            sections = ref_result.sections

        doc = Document()

        # Configure page layout
        self._setup_page(doc)

        # Title page
        self._add_title_page(doc, outline.title, university, discipline, author)

        # Table of contents placeholder
        self._add_toc_placeholder(doc)

        # Introduction
        intro_sections = [s for s in sections if s.section_title == "Введение"]
        if intro_sections:
            self._add_heading(doc, "ВВЕДЕНИЕ", level=1, numbered=False)
            intro_text = self._strip_leading_heading(intro_sections[0].content, "ВВЕДЕНИЕ")
            self._add_body_text(doc, intro_text)
            doc.add_page_break()

        # Chapters
        for chapter in outline.chapters:
            chapter_sections = [
                s for s in sections if s.chapter_number == chapter.number
            ]

            # Chapter heading
            self._add_heading(
                doc,
                f"{chapter.number} {chapter.title}",
                level=1,
            )

            for section in chapter_sections:
                # Section heading
                self._add_heading(doc, section.section_title, level=2)
                # Section body (strip duplicate heading if present)
                body = self._strip_leading_heading(section.content, section.section_title)
                self._add_body_text(doc, body)

            doc.add_page_break()

        # Conclusion
        conclusion_sections = [s for s in sections if s.section_title == "Заключение"]
        if conclusion_sections:
            self._add_heading(doc, "ЗАКЛЮЧЕНИЕ", level=1, numbered=False)
            concl_text = self._strip_leading_heading(
                conclusion_sections[0].content, "ЗАКЛЮЧЕНИЕ"
            )
            self._add_body_text(doc, concl_text)
            doc.add_page_break()

        # Bibliography — prefer registry (real sources), fall back to research sources
        self._add_bibliography(doc, sources, bibliography=bibliography)

        # Serialize to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        doc_bytes = buffer.getvalue()
        logger.info("docx_generated", size_kb=len(doc_bytes) // 1024)
        return doc_bytes

    def _setup_page(self, doc: Document) -> None:
        """Configure page size and margins per ГОСТ template."""
        t = self._template
        section = doc.sections[0]
        section.page_width = Mm(t.page_width_mm)
        section.page_height = Mm(t.page_height_mm)
        section.top_margin = Mm(t.margins.top_mm)
        section.bottom_margin = Mm(t.margins.bottom_mm)
        section.left_margin = Mm(t.margins.left_mm)
        section.right_margin = Mm(t.margins.right_mm)

    def _add_title_page(
        self,
        doc: Document,
        title: str,
        university: str,
        discipline: str,
        author: str,
    ) -> None:
        """Add a standard ГОСТ title page."""
        t = self._template

        # University name at top
        if university:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(university.upper())
            run.font.name = t.body.font.name
            run.font.size = Pt(14)
            run.font.bold = True

        # Spacer
        for _ in range(4):
            doc.add_paragraph()

        # "КУРСОВАЯ РАБОТА" label
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("КУРСОВАЯ РАБОТА")
        run.font.name = t.body.font.name
        run.font.size = Pt(18)
        run.font.bold = True

        # Discipline
        if discipline:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            clean_discipline = self._strip_quotes(discipline)
            run = p.add_run(f"по дисциплине «{clean_discipline}»")
            run.font.name = t.body.font.name
            run.font.size = Pt(14)

        # Spacer
        doc.add_paragraph()

        # Title
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        clean_title = self._strip_quotes(title)
        run = p.add_run(f"на тему: «{clean_title}»")
        run.font.name = t.body.font.name
        run.font.size = Pt(16)
        run.font.bold = True

        # Spacer
        for _ in range(6):
            doc.add_paragraph()

        # Author info (right-aligned)
        if author:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(f"Выполнил(а): {author}")
            run.font.name = t.body.font.name
            run.font.size = Pt(14)

        doc.add_page_break()

    def _add_toc_placeholder(self, doc: Document) -> None:
        """Add a table of contents placeholder.

        Note: python-docx cannot generate automatic TOC — it requires Word
        to update fields. We add a placeholder that can be updated when
        the document is opened in Word.
        """
        self._add_heading(doc, "СОДЕРЖАНИЕ", level=1, numbered=False)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("[Обновите оглавление: ПКМ → Обновить поле]")
        run.font.name = self._template.body.font.name
        run.font.size = Pt(12)
        run.font.italic = True
        run.font.color.rgb = RGBColor(128, 128, 128)

        doc.add_page_break()

    def _add_heading(
        self,
        doc: Document,
        text: str,
        level: int = 1,
        numbered: bool = True,
    ) -> None:
        """Add a formatted heading."""
        t = self._template
        style_config = getattr(t, f"heading_{level}", t.heading_1)

        heading = doc.add_heading(text, level=level)
        heading.alignment = ALIGNMENT_MAP.get(
            style_config.alignment, WD_ALIGN_PARAGRAPH.CENTER
        )

        # Format the heading run
        for run in heading.runs:
            run.font.name = style_config.font.name
            run.font.size = Pt(style_config.font.size_pt)
            run.font.bold = style_config.font.bold
            run.font.color.rgb = RGBColor(0, 0, 0)

        # Spacing
        pf = heading.paragraph_format
        pf.space_before = Pt(style_config.space_before_pt)
        pf.space_after = Pt(style_config.space_after_pt)

    @staticmethod
    def _clean_text(text: str) -> str:
        """Remove markdown formatting and HTML entities from LLM output."""
        # Remove HTML entities
        text = text.replace("&nbsp;", " ")
        text = re.sub(r"&[a-zA-Z]+;", " ", text)
        # Remove markdown headings (## Title -> Title)
        text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)
        # Remove bold **text** -> text
        text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
        # Remove italic *text* -> text (but not citation [*])
        text = re.sub(r"(?<!\[)\*(.+?)\*(?!\])", r"\1", text)
        # Remove inline code `text` -> text
        text = re.sub(r"`(.+?)`", r"\1", text)
        # Remove markdown links [text](url) -> text
        text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
        # Remove bullet points (- item, * item)
        text = re.sub(r"^\s*[-*•]\s+", "", text, flags=re.MULTILINE)
        # Remove numbered list prefixes (1. item, 1) item) — but not citation [1]
        text = re.sub(r"^\s*\d+[.)]\s+", "", text, flags=re.MULTILINE)
        # Collapse multiple spaces
        text = re.sub(r"  +", " ", text)
        return text

    def _add_body_text(self, doc: Document, text: str) -> None:
        """Add body text with proper ГОСТ formatting.

        Strips markdown/HTML artifacts and splits into paragraphs.
        """
        t = self._template
        body_style = t.body

        text = self._clean_text(text)

        paragraphs = text.split("\n")
        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text:
                continue

            p = doc.add_paragraph()
            p.alignment = ALIGNMENT_MAP.get(
                body_style.alignment, WD_ALIGN_PARAGRAPH.JUSTIFY
            )

            # Paragraph formatting
            pf = p.paragraph_format
            pf.first_line_indent = Mm(body_style.first_line_indent_mm)
            pf.space_after = Pt(body_style.space_after_pt)
            pf.line_spacing = body_style.line_spacing

            # Handle citation references [N] with different formatting
            parts = re.split(r"(\[\d+\])", para_text)
            for part in parts:
                if re.match(r"\[\d+\]", part):
                    run = p.add_run(part)
                    run.font.name = body_style.font.name
                    run.font.size = Pt(body_style.font.size_pt)
                    run.font.bold = False
                else:
                    run = p.add_run(part)
                    run.font.name = body_style.font.name
                    run.font.size = Pt(body_style.font.size_pt)

    @staticmethod
    def _strip_quotes(text: str) -> str:
        """Strip surrounding «» and "" quotes from user input.

        Prevents double quoting: user enters «Менеджмент», template wraps in «»,
        result would be ««Менеджмент»» without this.
        """
        text = text.strip()
        if text.startswith("«") and text.endswith("»"):
            text = text[1:-1].strip()
        if text.startswith('"') and text.endswith('"'):
            text = text[1:-1].strip()
        return text

    @staticmethod
    def _strip_leading_heading(text: str, heading: str) -> str:
        """Remove a duplicate heading from the start of LLM-generated text.

        Handles various formats:
        - Exact heading match (case-insensitive): "ВВЕДЕНИЕ", "1.1 Title"
        - "РАЗДЕЛ: 1.1 Title" prefix
        - Numbered patterns: "1.1 Title" or "1.1. Title"
        """
        stripped = text.lstrip()
        if not stripped:
            return text
        heading_upper = heading.upper().strip()

        # Check for exact heading match (case-insensitive)
        if stripped.upper().startswith(heading_upper):
            after = stripped[len(heading_upper):]
            after = after.lstrip(" \t\n\r.:;-—")
            if after:
                return after

        # Check for "РАЗДЕЛ: <heading>" prefix
        razdel_match = re.match(
            r"^\s*(?:РАЗДЕЛ|Раздел)\s*:?\s*", stripped, re.IGNORECASE
        )
        if razdel_match:
            after_razdel = stripped[razdel_match.end():]
            if after_razdel.upper().startswith(heading_upper):
                after = after_razdel[len(heading_upper):]
                after = after.lstrip(" \t\n\r.:;-—")
                if after:
                    return after
            # Even if heading doesn't match exactly, strip the whole first line
            first_newline = stripped.find("\n")
            if first_newline > 0:
                after = stripped[first_newline:].lstrip("\n\r")
                if after:
                    return after

        return text

    def _add_bibliography(
        self,
        doc: Document,
        sources: list[Source],
        bibliography: BibliographyRegistry | None = None,
    ) -> None:
        """Add bibliography / references section.

        Uses the bibliography registry (built from real research sources) when
        available. Falls back to formatting raw Source objects.
        Skips the entire section if there are no entries to render.
        """
        has_registry_entries = bibliography is not None and len(bibliography.entries) > 0
        has_sources = len(sources) > 0

        logger.info(
            "adding_bibliography",
            has_registry=bibliography is not None,
            registry_entries=len(bibliography.entries) if bibliography else 0,
            fallback_sources=len(sources),
        )

        # Skip entire bibliography section if nothing to render
        if not has_registry_entries and not has_sources:
            logger.warning(
                "bibliography_empty_skipping_section",
                reason="No bibliography entries and no fallback sources available",
            )
            return

        self._add_heading(doc, "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", level=1, numbered=False)

        # Prefer bibliography registry (real, verified sources)
        if has_registry_entries:
            for entry in bibliography.entries:  # type: ignore[union-attr]
                self._add_bib_entry(doc, f"{entry.number}. {entry.formatted_reference}")
            return

        # Fallback to research sources (same format as registry would produce)
        for i, source in enumerate(sources, 1):
            bib_entry = f"{i}. {source.title}"
            if source.url:
                bib_entry += f" [Электронный ресурс]. — URL: {source.url}"
            self._add_bib_entry(doc, bib_entry)

    def _add_bib_entry(self, doc: Document, text: str) -> None:
        """Add a single bibliography entry with proper formatting."""
        t = self._template
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf = p.paragraph_format
        pf.first_line_indent = Mm(t.body.first_line_indent_mm)
        pf.line_spacing = t.body.line_spacing

        run = p.add_run(text)
        run.font.name = t.body.font.name
        run.font.size = Pt(t.body.font.size_pt)
