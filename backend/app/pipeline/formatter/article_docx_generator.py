"""ГОСТ-compliant .docx generator for scientific articles."""

import io
import re

import structlog
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm, Pt, RGBColor

from backend.app.pipeline.formatter.reference_extractor import (
    extract_and_renumber_references,
)
from shared.schemas.pipeline import (
    FactCheckResult,
    Outline,
    SectionContent,
    Source,
)
from shared.schemas.template import GostTemplate

logger = structlog.get_logger()

ALIGNMENT_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


class ArticleDocxGenerator:
    """Generates ГОСТ-compliant .docx files for scientific articles.

    Structure: Title → Abstract → Keywords → Introduction →
    Sections → Conclusion → References
    """

    def __init__(self, template: GostTemplate | None = None) -> None:
        self._template = template or GostTemplate()

    def generate(
        self,
        outline: Outline,
        sections: list[SectionContent],
        sources: list[Source],
        fact_check: FactCheckResult | None = None,
        university: str = "",
        discipline: str = "",
        author: str = "",
    ) -> bytes:
        """Generate a complete article .docx document."""
        # Extract inline references and build unified bibliography
        ref_result = extract_and_renumber_references(sections)
        sections = ref_result.sections
        collected_bibliography = ref_result.bibliography

        doc = Document()
        t = self._template

        self._setup_page(doc)

        # Article title (centered, bold)
        self._add_article_title(doc, outline.title)

        # Author info
        if author:
            self._add_author_info(doc, author, university)

        # Abstract
        abstract_sections = [s for s in sections if s.section_title == "Аннотация"]
        if abstract_sections:
            self._add_labeled_section(doc, "Аннотация", abstract_sections[0].content)

        # Keywords
        if outline.keywords:
            self._add_keywords(doc, outline.keywords)

        # Introduction
        intro_sections = [s for s in sections if s.section_title == "Введение"]
        if intro_sections:
            self._add_heading(doc, "Введение", level=1)
            self._add_body_text(doc, intro_sections[0].content)

        # Main sections (flat structure, no subsections)
        for chapter in outline.chapters:
            chapter_sections = [
                s for s in sections if s.chapter_number == chapter.number
            ]

            self._add_heading(doc, f"{chapter.number}. {chapter.title}", level=1)

            for section in chapter_sections:
                self._add_body_text(doc, section.content)

        # Conclusion
        conclusion_sections = [s for s in sections if s.section_title == "Заключение"]
        if conclusion_sections:
            self._add_heading(doc, "Заключение", level=1)
            self._add_body_text(doc, conclusion_sections[0].content)

        # References — prefer extracted inline refs, fall back to research sources
        self._add_bibliography(doc, sources, collected_bibliography)

        # Serialize
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        doc_bytes = buffer.getvalue()
        logger.info("article_docx_generated", size_kb=len(doc_bytes) // 1024)
        return doc_bytes

    def _setup_page(self, doc: Document) -> None:
        t = self._template
        section = doc.sections[0]
        section.page_width = Mm(t.page_width_mm)
        section.page_height = Mm(t.page_height_mm)
        section.top_margin = Mm(t.margins.top_mm)
        section.bottom_margin = Mm(t.margins.bottom_mm)
        section.left_margin = Mm(t.margins.left_mm)
        section.right_margin = Mm(t.margins.right_mm)

    def _add_article_title(self, doc: Document, title: str) -> None:
        """Add article title — centered, bold, larger font."""
        t = self._template

        # Spacer
        doc.add_paragraph()

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.space_after = Pt(12)

        run = p.add_run(title.upper())
        run.font.name = t.body.font.name
        run.font.size = Pt(16)
        run.font.bold = True

    def _add_author_info(
        self, doc: Document, author: str, university: str
    ) -> None:
        t = self._template

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(author)
        run.font.name = t.body.font.name
        run.font.size = Pt(12)
        run.font.italic = True

        if university:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(university)
            run.font.name = t.body.font.name
            run.font.size = Pt(11)
            run.font.italic = True

        doc.add_paragraph()  # Spacer

    def _add_labeled_section(
        self, doc: Document, label: str, content: str
    ) -> None:
        """Add a labeled section (e.g., Abstract) with bold label inline."""
        t = self._template
        content = self._clean_text(content)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf = p.paragraph_format
        pf.first_line_indent = Mm(t.body.first_line_indent_mm)
        pf.line_spacing = t.body.line_spacing
        pf.space_after = Pt(6)

        # Bold label
        run = p.add_run(f"{label}. ")
        run.font.name = t.body.font.name
        run.font.size = Pt(t.body.font.size_pt)
        run.font.bold = True

        # Content text
        run = p.add_run(content)
        run.font.name = t.body.font.name
        run.font.size = Pt(t.body.font.size_pt)

    def _add_keywords(self, doc: Document, keywords: list[str]) -> None:
        t = self._template

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf = p.paragraph_format
        pf.first_line_indent = Mm(t.body.first_line_indent_mm)
        pf.space_after = Pt(12)

        run = p.add_run("Ключевые слова: ")
        run.font.name = t.body.font.name
        run.font.size = Pt(t.body.font.size_pt)
        run.font.bold = True

        run = p.add_run(", ".join(keywords) + ".")
        run.font.name = t.body.font.name
        run.font.size = Pt(t.body.font.size_pt)
        run.font.italic = True

    def _add_heading(
        self, doc: Document, text: str, level: int = 1
    ) -> None:
        t = self._template
        style_config = getattr(t, f"heading_{level}", t.heading_1)

        heading = doc.add_heading(text, level=level)
        heading.alignment = ALIGNMENT_MAP.get(
            style_config.alignment, WD_ALIGN_PARAGRAPH.CENTER
        )

        for run in heading.runs:
            run.font.name = style_config.font.name
            run.font.size = Pt(style_config.font.size_pt)
            run.font.bold = style_config.font.bold
            run.font.color.rgb = RGBColor(0, 0, 0)

        pf = heading.paragraph_format
        pf.space_before = Pt(style_config.space_before_pt)
        pf.space_after = Pt(style_config.space_after_pt)

    @staticmethod
    def _clean_text(text: str) -> str:
        """Remove markdown formatting and HTML entities from LLM output."""
        text = text.replace("&nbsp;", " ")
        text = re.sub(r"&[a-zA-Z]+;", " ", text)
        text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)
        text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
        text = re.sub(r"(?<!\[)\*(.+?)\*(?!\])", r"\1", text)
        text = re.sub(r"`(.+?)`", r"\1", text)
        text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
        text = re.sub(r"^\s*[-*•]\s+", "", text, flags=re.MULTILINE)
        text = re.sub(r"^\s*\d+[.)]\s+", "", text, flags=re.MULTILINE)
        text = re.sub(r"  +", " ", text)
        return text

    def _add_body_text(self, doc: Document, text: str) -> None:
        t = self._template
        body_style = t.body

        text = self._clean_text(text)

        paragraphs = text.split("\n")
        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text:
                continue

            p = doc.add_paragraph()
            p.alignment = ALIGNMENT_MAP.get(
                body_style.alignment, WD_ALIGN_PARAGRAPH.JUSTIFY
            )

            pf = p.paragraph_format
            pf.first_line_indent = Mm(body_style.first_line_indent_mm)
            pf.space_after = Pt(body_style.space_after_pt)
            pf.line_spacing = body_style.line_spacing

            parts = re.split(r"(\[\d+\])", para_text)
            for part in parts:
                if re.match(r"\[\d+\]", part):
                    run = p.add_run(part)
                    run.font.name = body_style.font.name
                    run.font.size = Pt(body_style.font.size_pt)
                    run.font.bold = False
                else:
                    run = p.add_run(part)
                    run.font.name = body_style.font.name
                    run.font.size = Pt(body_style.font.size_pt)

    def _add_bibliography(
        self,
        doc: Document,
        sources: list[Source],
        collected_refs: list[str] | None = None,
    ) -> None:
        """Add bibliography section.

        Uses references extracted from LLM text when available,
        falls back to research sources otherwise.
        """
        self._add_heading(doc, "СПИСОК ЛИТЕРАТУРЫ", level=1)

        t = self._template

        if collected_refs:
            for i, ref_text in enumerate(collected_refs, 1):
                self._add_bib_entry(doc, f"{i}. {ref_text}")
            return

        for i, source in enumerate(sources, 1):
            bib_entry = f"{i}. {source.title}"
            if source.url:
                bib_entry += f" [Электронный ресурс]. — URL: {source.url}"
            self._add_bib_entry(doc, bib_entry)

    def _add_bib_entry(self, doc: Document, text: str) -> None:
        t = self._template
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf = p.paragraph_format
        pf.first_line_indent = Mm(t.body.first_line_indent_mm)
        pf.line_spacing = t.body.line_spacing

        run = p.add_run(text)
        run.font.name = t.body.font.name
        run.font.size = Pt(t.body.font.size_pt)
