"""Load or generate a GOST reference .docx for visual template matching.

Priority:
1. User-provided file from templates/ directory
2. Programmatically generated fallback
"""

import io
from pathlib import Path

import structlog
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm, Pt

from shared.schemas.template import GostTemplate

logger = structlog.get_logger()

# Project root: 5 levels up from this file
# backend/app/pipeline/formatter/gost_reference.py → backend/app/pipeline/formatter → project root
_PROJECT_ROOT = Path(__file__).resolve().parent.parent.parent.parent.parent

_DEFAULT_TEMPLATE_PATH = _PROJECT_ROOT / "templates" / "tamplate_course_work.docx"


def get_default_reference() -> bytes:
    """Load the default GOST reference .docx.

    Tries to load from templates/ directory first,
    falls back to programmatic generation.
    """
    if _DEFAULT_TEMPLATE_PATH.exists():
        logger.info("loading_reference_template", path=str(_DEFAULT_TEMPLATE_PATH))
        return _DEFAULT_TEMPLATE_PATH.read_bytes()

    logger.info("generating_fallback_reference")
    return _generate_fallback_reference()


def _generate_fallback_reference(template: GostTemplate | None = None) -> bytes:
    """Generate a minimal GOST reference .docx programmatically."""
    t = template or GostTemplate()
    doc = Document()

    section = doc.sections[0]
    section.page_width = Mm(t.page_width_mm)
    section.page_height = Mm(t.page_height_mm)
    section.top_margin = Mm(t.margins.top_mm)
    section.bottom_margin = Mm(t.margins.bottom_mm)
    section.left_margin = Mm(t.margins.left_mm)
    section.right_margin = Mm(t.margins.right_mm)

    font_name = t.body.font.name
    font_size = Pt(t.body.font.size_pt)
    line_spacing = t.body.line_spacing
    indent = Mm(t.body.first_line_indent_mm)

    def add_heading_text(text: str, level: int = 1) -> None:
        style_config = getattr(t, f"heading_{level}", t.heading_1)
        heading = doc.add_heading(text, level=level)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in heading.runs:
            run.font.name = style_config.font.name
            run.font.size = Pt(style_config.font.size_pt)
            run.font.bold = style_config.font.bold
        pf = heading.paragraph_format
        pf.space_before = Pt(style_config.space_before_pt)
        pf.space_after = Pt(style_config.space_after_pt)

    def add_body(text: str) -> None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf = p.paragraph_format
        pf.first_line_indent = indent
        pf.line_spacing = line_spacing
        pf.space_after = Pt(t.body.space_after_pt)
        run = p.add_run(text)
        run.font.name = font_name
        run.font.size = font_size

    # Title page
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("КУРСОВАЯ РАБОТА")
    run.font.name = font_name
    run.font.size = Pt(16)
    run.font.bold = True

    # Введение
    doc.add_page_break()
    add_heading_text("ВВЕДЕНИЕ", level=1)
    add_body(
        "Актуальность темы исследования обусловлена возрастающей ролью цифровых "
        "технологий в современной экономике. Процессы цифровой трансформации "
        "затрагивают все сферы деятельности организаций."
    )

    # Глава 1
    doc.add_page_break()
    add_heading_text("ГЛАВА 1. ТЕОРЕТИЧЕСКИЕ ОСНОВЫ", level=1)
    h2 = doc.add_heading("1.1 Основные понятия", level=2)
    h2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h2.runs:
        run.font.name = font_name
        run.font.size = Pt(t.heading_2.font.size_pt)
        run.font.bold = True
    add_body(
        "Согласно исследованиям отечественных авторов, цифровизация является "
        "ключевым фактором повышения конкурентоспособности предприятий [1]. "
        "По данным Росстата, более 60% крупных компаний активно внедряют "
        "элементы цифровой трансформации [2]."
    )

    # Заключение
    doc.add_page_break()
    add_heading_text("ЗАКЛЮЧЕНИЕ", level=1)
    add_body(
        "В результате проведённого исследования были достигнуты поставленные "
        "цели и решены задачи, сформулированные во введении."
    )

    # Список литературы
    doc.add_page_break()
    add_heading_text("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", level=1)
    for i, ref in enumerate(["Глазьев С.Ю. Великая цифровая революция. — М., 2023.", "Росстат. Индикаторы цифровой экономики. — М., 2024."], 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf = p.paragraph_format
        pf.first_line_indent = indent
        pf.line_spacing = line_spacing
        run = p.add_run(f"{i}. {ref}")
        run.font.name = font_name
        run.font.size = font_size

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
